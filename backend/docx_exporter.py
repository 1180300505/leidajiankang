from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import datetime, timedelta
import os

# 健康度/分数/风险等级对应的表格底色（十六进制，无#）
LEVEL_FILL = {
    "优秀": "90EE90",   # 浅绿色
    "注意": "FFFF99",   # 浅黄色
    "严重": "FF6B6B",   # 浅红色
    "危险": "FF6B6B",   # 浅红色（等同于严重）
    "高": "FF6B6B",     # 风险等级高 → 红色
    "中": "FFFF99",     # 风险等级中 → 黄色
}

def _shade_cell(cell, fill_hex):
    """为表格单元格设置背景底色。fill_hex: 6位十六进制色值，如 'FFFF99'"""
    if not fill_hex:
        return
    shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill_hex))
    cell._tc.get_or_add_tcPr().append(shd)

def _level_to_fill(level):
    """根据健康等级返回对应的底色。level: 优秀/注意/严重/危险/良好"""
    return LEVEL_FILL.get(level, None)

def _score_to_fill(score, grade="正常"):
    """根据分数和等级返回对应的底色"""
    level = _score_grade_to_level(score, grade)
    return _level_to_fill(level)


def _avg(items, key, default=0):
    vals = [r.get(key) for r in items if r.get(key) is not None]
    try:
        return round(sum(float(x) for x in vals) / len(vals), 1) if vals else default
    except (TypeError, ValueError):
        return default


def _score_grade_to_level(score, grade):
    if score >= 90:
        return "优秀"
    if score >= 70:
        return "注意" if grade == "轻度异常" else "良好"
    return "危险"


def _score_to_display_level(score):
    if score >= 90:
        return "优秀"
    if score >= 70:
        return "良好"
    return "危险"


def _score_to_rul(score):
    if score >= 85:
        return (450, 400, 500, "中")
    if score >= 70:
        return (320, 280, 360, "高")
    return (180, 150, 210, "高")


def _score_to_maintenance(score, name):
    if score < 70:
        return ("立即执行", [f"检查{name}相关部件（电机、接线、冷却）"])
    if score < 85:
        return ("计划性维护", [f"1-3个月内安排{name}检查或更换"])
    return ("例行巡检", [f"{name}状态良好，保持常规巡检"])

def export_full_docx(data_list, start_time, end_time, filename="detailed_report.docx"):
    """
    适配后端的导出函数
    :param data_list: 数据库查询出的字典列表
    :param start_time: 前端传来的开始时间字符串
    :param end_time: 前端传来的结束时间字符串
    :param filename: 生成的临时文件名
    """
    doc = Document()
    
    # 1. 标题与范围声明
    title = doc.add_heading('设备全参数运行日志详细报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run(f"统计周期：{start_time} 至 {end_time}")
    run.font.size = Pt(11)
    
    if not data_list:
        doc.add_paragraph("\n查询时间范围内无记录。")
        doc.save(filename)
        return filename

    # 2. 定义字段逻辑分组 (可根据实际需求增减)
    sections = {
        "1. 基本信息与系统状态": [
            ('timestamp', '记录时间'), ('sys_mode', '工作模式'), 
            ('sys_signal_source', '信号源'), ('sys_lock_status', '锁定状态'),
            ('sys_lock_indicator', '锁定指示')
        ],
        "2. 信号详细参数": [
            ('sig_agc_threshold', 'AGC门限'), ('sig_agc_voltage', 'AGC电压'),
            ('sig_azimuth_err_v', '方位误差电压'), ('sig_pitch_err_v', '俯仰误差电压')
        ],
        "3. 转台系跟踪数据": [
            ('tt_guide_az', '引导方位'), ('tt_curr_az', '当前方位'), ('tt_dev_az', '方位偏差'),
            ('tt_guide_pt', '引导俯仰'), ('tt_curr_pt', '当前俯仰'), ('tt_dev_pt', '俯仰偏差')
        ],
        "4. 电机诊断数据": [
            ('m1_status', '电机1状态'), ('m1_temp', '电机1温度'), ('m1_current', '电机1电流'),
            ('m2_status', '电机2状态'), ('m2_temp', '电机2温度'), ('m2_current', '电机2电流')
        ]
    }

    # 3. 循环生成内容
    for index, record in enumerate(data_list):
        # 每一条数据的大标题
        doc.add_heading(f"记录 #{index + 1} | 录入时间: {record.get('timestamp')}", level=1)
        
        for section_name, fields in sections.items():
            doc.add_heading(section_name, level=2)
            
            # 创建 2 列的窄表（Key-Value 模式）
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            table.autofit = False 
            
            for key, label in fields:
                row_cells = table.add_row().cells
                # 第一列：参数名（加粗，固定宽度）
                row_cells[0].width = Inches(1.8)
                row_cells[0].text = label
                row_cells[0].paragraphs[0].runs[0].font.bold = True
                
                # 第二列：具体数值
                raw_val = record.get(key)
                val = str(raw_val) if raw_val is not None else "N/A"
                row_cells[1].text = val
                
                # 针对温度的阈值报警（红色高亮）
                if "temp" in key:
                    try:
                        if float(val) > 50:
                            run = row_cells[1].paragraphs[0].runs[0]
                            run.font.color.rgb = RGBColor(200, 0, 0)
                            run.font.bold = True
                    except:
                        pass
            
            doc.add_paragraph() # 段落间距

        # 每条记录占一页或多页，最后一条不加分页符
        if index < len(data_list) - 1:
            doc.add_page_break()

    # 4. 保存
    doc.save(filename)
    return filename


def export_health_report_docx(health_records, system_logs, start_time, end_time, algorithm, prev_records=None, filename=None):
    """
    导出健康评估报告（Word）
    :param health_records: 该算法在时间段内的健康记录
    :param system_logs: 同一时间段的原始日志（部件描述用）
    :param start_time: 报告开始时间
    :param end_time: 报告结束时间
    :param algorithm: 算法名 kmeans/som
    :param prev_records: 上一周期健康记录（用于对比）
    :param filename: 输出文件名
    """
    prev_records = prev_records or []
    if filename is None:
        _dir = os.path.dirname(os.path.abspath(__file__))
        filename = os.path.join(_dir, "health_report.docx")

    doc = Document()
    algo_label = "KMeans" if algorithm == "kmeans" else "SOM"

    # 1. 标题与周期
    title = doc.add_heading("天线系统健康评估报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").font.size = Pt(10)
    doc.add_paragraph()
    period = doc.add_paragraph()
    period.add_run(f"统计周期：{start_time} 至 {end_time}").font.bold = True
    period.add_run(f"  |  评估算法：{algo_label}")
    doc.add_paragraph()

    # 2. 整体健康状态概览
    avg_overall = _avg(health_records, "overall_score", 85)
    avg_prev = _avg(prev_records, "overall_score") if prev_records else avg_overall
    diff = round(avg_overall - avg_prev, 1)
    if diff > 0:
        diff_str = f"↑ +{diff}"
    elif diff < 0:
        diff_str = f"↓ {diff}"
    else:
        diff_str = "→ 持平"
    level = _score_grade_to_level(avg_overall, health_records[-1].get("overall_grade", "正常") if health_records else "正常")

    doc.add_heading("一、整体健康状态概览", level=1)
    table1 = doc.add_table(rows=2, cols=4)
    table1.style = "Table Grid"
    hdr = table1.rows[0].cells
    hdr[0].text = "评估项"
    hdr[1].text = "分数"
    hdr[2].text = "健康等级"
    hdr[3].text = "与上周期对比"
    for c in hdr:
        c.paragraphs[0].runs[0].font.bold = True
    row1 = table1.rows[1].cells
    row1[0].text = "系统综合健康"
    row1[1].text = str(avg_overall)
    row1[2].text = level
    row1[3].text = diff_str
    fill1 = _score_to_fill(avg_overall, health_records[-1].get("overall_grade", "正常") if health_records else "正常")
    if fill1:
        _shade_cell(row1[1], fill1)
        _shade_cell(row1[2], fill1)
    doc.add_paragraph()

    comment = "系统本月运行总体平稳。"
    if avg_overall < 70:
        comment += "存在明显异常，建议尽快检查相关子系统。"
    elif avg_overall < 85:
        comment += "部分子系统有轻微退化迹象，建议加强监测。"
    doc.add_paragraph(f"整体评语：{comment}")
    doc.add_paragraph()

    # 3. 分层级健康度分析
    doc.add_heading("二、分层级健康度分析", level=1)
    avg_tt = _avg(health_records, "turntable_score", 85)
    avg_ef = _avg(health_records, "electrofeed_score", 85)
    level_tt = _score_grade_to_level(avg_tt, health_records[-1].get("turntable_grade", "正常") if health_records else "正常")
    level_ef = _score_grade_to_level(avg_ef, health_records[-1].get("electrofeed_grade", "正常") if health_records else "正常")

    table2 = doc.add_table(rows=3, cols=4)
    table2.style = "Table Grid"
    h2 = table2.rows[0].cells
    h2[0].text = "子系统"
    h2[1].text = "健康分数"
    h2[2].text = "健康等级"
    h2[3].text = "关键指标"
    for c in h2:
        c.paragraphs[0].runs[0].font.bold = True
    rtt = table2.rows[1].cells
    rtt[0].text = "转台系"
    rtt[1].text = str(avg_tt)
    rtt[2].text = level_tt
    rtt[3].text = "方位/俯仰角度偏差在阈值内" if avg_tt >= 70 else "方位/俯仰偏差需关注"
    fill_tt = _level_to_fill(level_tt)
    if fill_tt:
        _shade_cell(rtt[1], fill_tt)
        _shade_cell(rtt[2], fill_tt)
    ref = table2.rows[2].cells
    ref[0].text = "电馈系"
    ref[1].text = str(avg_ef)
    ref[2].text = level_ef
    ref[3].text = "电机温度、电压正常" if avg_ef >= 70 else "电机参数波动需关注"
    fill_ef = _level_to_fill(level_ef)
    if fill_ef:
        _shade_cell(ref[1], fill_ef)
        _shade_cell(ref[2], fill_ef)
    doc.add_paragraph()

    # 4. 部件健康度分析
    doc.add_heading("三、部件健康度分析", level=1)
    m1_score = avg_ef
    m2_score = avg_ef
    m1_feature = "温度、电流正常"
    m2_feature = "温度、电流正常"
    if system_logs:
        try:
            temps = [(float(r.get("m1_temp") or 0), float(r.get("m2_temp") or 0)) for r in system_logs]
            if temps:
                t1_avg = sum(t[0] for t in temps) / len(temps)
                t2_avg = sum(t[1] for t in temps) / len(temps)
                if t1_avg > 50:
                    m1_feature = "绕组温度梯度异常"
                    m1_score = min(m1_score, 76)
                if t2_avg > 50:
                    m2_feature = "绕组温度梯度异常"
                    m2_score = min(m2_score, 76)
        except (TypeError, ValueError):
            pass

    table3 = doc.add_table(rows=3, cols=4)
    table3.style = "Table Grid"
    h3 = table3.rows[0].cells
    h3[0].text = "部件"
    h3[1].text = "部件类型"
    h3[2].text = "健康分数"
    h3[3].text = "主要退化特征/建议"
    for c in h3:
        c.paragraphs[0].runs[0].font.bold = True
    rm1 = table3.rows[1].cells
    rm1[0].text = "电馈系 - 电机1"
    rm1[1].text = "永磁同步电机"
    rm1[2].text = str(round(m1_score, 1))
    rm1[3].text = m1_feature + ("；3个月内更换" if m1_score < 70 else ("；6个月后检查" if m1_score < 85 else "；例行巡检"))
    fill_m1 = _score_to_fill(m1_score)
    if fill_m1:
        _shade_cell(rm1[2], fill_m1)
    rm2 = table3.rows[2].cells
    rm2[0].text = "电馈系 - 电机2"
    rm2[1].text = "永磁同步电机"
    rm2[2].text = str(round(m2_score, 1))
    rm2[3].text = m2_feature + ("；3个月内更换" if m2_score < 70 else ("；6个月后检查" if m2_score < 85 else "；例行巡检"))
    fill_m2 = _score_to_fill(m2_score)
    if fill_m2:
        _shade_cell(rm2[2], fill_m2)
    doc.add_paragraph()

    # 5. 寿命预测与风险预警
    doc.add_heading("四、寿命预测与风险预警", level=1)
    rul_overall = _score_to_rul(avg_overall)
    rul_tt = _score_to_rul(avg_tt)
    rul_ef = _score_to_rul(avg_ef)

    table4 = doc.add_table(rows=4, cols=4)
    table4.style = "Table Grid"
    h4 = table4.rows[0].cells
    h4[0].text = "关键部件"
    h4[1].text = "当前剩余使用寿命(RUL)"
    h4[2].text = "预测失效时间窗口"
    h4[3].text = "风险等级"
    for c in h4:
        c.paragraphs[0].runs[0].font.bold = True
    base = datetime.now()
    r1 = table4.rows[1].cells
    r1[0].text = "转台系"
    r1[1].text = f"{rul_tt[0]}天 ({rul_tt[1]}-{rul_tt[2]}天)"
    fail_start = (base + timedelta(days=rul_tt[1])).strftime("%Y-%m")
    fail_end = (base + timedelta(days=rul_tt[2])).strftime("%Y-%m")
    r1[2].text = f"{fail_start} 至 {fail_end}"
    r1[3].text = rul_tt[3]
    fill_r1 = _level_to_fill(rul_tt[3])
    if fill_r1:
        _shade_cell(r1[3], fill_r1)
    r2 = table4.rows[2].cells
    r2[0].text = "电馈系 - 电机1"
    r2[1].text = f"{rul_ef[0]}天 ({rul_ef[1]}-{rul_ef[2]}天)"
    r2[2].text = f"{(base + timedelta(days=rul_ef[1])).strftime('%Y-%m')} 至 {(base + timedelta(days=rul_ef[2])).strftime('%Y-%m')}"
    r2[3].text = rul_ef[3]
    fill_r2 = _level_to_fill(rul_ef[3])
    if fill_r2:
        _shade_cell(r2[3], fill_r2)
    r3 = table4.rows[3].cells
    r3[0].text = "电馈系 - 电机2"
    r3[1].text = f"{rul_ef[0]}天 ({rul_ef[1]}-{rul_ef[2]}天)"
    r3[2].text = f"{(base + timedelta(days=rul_ef[1])).strftime('%Y-%m')} 至 {(base + timedelta(days=rul_ef[2])).strftime('%Y-%m')}"
    r3[3].text = rul_ef[3]
    fill_r3 = _level_to_fill(rul_ef[3])
    if fill_r3:
        _shade_cell(r3[3], fill_r3)
    doc.add_paragraph()

    # 6. 维护建议
    doc.add_heading("五、维护建议", level=1)
    imm_items = []
    plan_items = []
    if avg_overall < 70 or avg_tt < 70:
        imm_items.append("检查转台系相关部件（方位/俯仰驱动、接线）")
    if avg_overall < 70 or avg_ef < 70:
        imm_items.append("检查电馈系电机1、电机2的接线端子与冷却风扇")
        imm_items.append("对电机进行油样/温度分析")
    if 70 <= avg_overall < 85 or 70 <= avg_ef < 85:
        plan_items.append("订购电馈系电机备件(物料号: MOT-EF-001)")
        plan_items.append("安排电机振动测试")
    if avg_overall >= 85:
        plan_items.append("保持例行巡检")

    if imm_items:
        doc.add_paragraph("立即执行（1周内）：", style="List Bullet").runs[0].font.bold = True
        for item in imm_items:
            doc.add_paragraph(item, style="List Bullet")
    if plan_items:
        doc.add_paragraph("计划性维护（1-3个月内）：", style="List Bullet").runs[0].font.bold = True
        for item in plan_items:
            doc.add_paragraph(item, style="List Bullet")
    if not imm_items and not plan_items:
        doc.add_paragraph("当前无紧急维护项，建议保持常规巡检。")

    doc.save(filename)
    return filename


def export_fault_report_docx(fault_report: dict, filename=None):
    """
    导出故障诊断报告（Word）
    :param fault_report: 故障报告字典，包含 event_id, 故障时间, 故障设备, 故障部件, 故障类型, 严重等级, 异常参数 等
    :param filename: 输出路径，可为 BytesIO
    """
    if filename is None:
        _dir = os.path.dirname(os.path.abspath(__file__))
        event_id = fault_report.get("event_id", "FAULT-UNKNOWN")
        safe_id = "".join(c if c.isalnum() or c in "-_" else "_" for c in event_id)
        filename = os.path.join(_dir, f"fault_report_{safe_id}.docx")

    doc = Document()

    # 1. 标题
    title = doc.add_heading("天线方位轴驱动子系统故障诊断报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 2. 事件 ID 与故障时间
    doc.add_paragraph(f"事件ID：{fault_report.get('event_id', 'N/A')}")
    doc.add_paragraph(f"故障时间：{fault_report.get('故障时间', 'N/A')}")
    doc.add_paragraph()

    # 3. 设备信息
    doc.add_heading("一、设备信息", level=1)
    table1 = doc.add_table(rows=4, cols=2)
    table1.style = "Table Grid"
    rows_data = [
        ("故障设备", fault_report.get("故障设备", "N/A")),
        ("故障部件", fault_report.get("故障部件", "N/A")),
        ("故障类型", fault_report.get("故障类型", "N/A")),
        ("严重等级", fault_report.get("严重等级", "N/A")),
    ]
    for i, (label, value) in enumerate(rows_data):
        table1.rows[i].cells[0].text = label
        table1.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        table1.rows[i].cells[1].text = str(value)
    doc.add_paragraph()

    # 4. 异常参数
    doc.add_heading("二、异常参数", level=1)
    params = fault_report.get("异常参数") or []
    if params:
        table2 = doc.add_table(rows=len(params) + 1, cols=4)
        table2.style = "Table Grid"
        hdr = table2.rows[0].cells
        hdr[0].text = "参数名称"
        hdr[1].text = "当前值"
        hdr[2].text = "阈值"
        hdr[3].text = "说明"
        for c in hdr:
            c.paragraphs[0].runs[0].font.bold = True
        for i, p in enumerate(params):
            r = table2.rows[i + 1].cells
            r[0].text = str(p.get("name", "N/A"))
            r[1].text = str(p.get("current", "N/A"))
            r[2].text = str(p.get("threshold", "N/A"))
            r[3].text = str(p.get("desc", ""))
    else:
        doc.add_paragraph("无异常参数记录。")
    doc.add_paragraph()

    # 5. 维护建议（缺省）
    doc.add_heading("三、维护建议", level=1)
    doc.add_paragraph("暂无具体维护建议，请结合实际工况进行排查。")
    doc.add_paragraph()

    # 6. 附件（缺省）
    doc.add_heading("四、附件", level=1)
    doc.add_paragraph("暂无附件。")

    doc.save(filename)
    return filename